import os
import uuid
import logging
import json
import cv2
import numpy as np
import re
import fitz  # PyMuPDF
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from PIL import Image
import io
from fuzzywuzzy import fuzz
from fastapi.responses import FileResponse

# Optional heavy imports (not needed for cloud deployment)
try:
    import torch
    from transformers import LayoutLMv3Processor, LayoutLMv3ForSequenceClassification
    TORCH_AVAILABLE = True
except ImportError:
    TORCH_AVAILABLE = False
    print("Warning: PyTorch/Transformers not available. Some features disabled.")

# Phase 2 Imports
from template_engine.template_models import TemplateSchema
from template_engine.template_extractor import extract_template_schema
from template_engine.template_manager import register_template, list_templates, get_template_schema
from template_engine.template_mapper import fill_template

# Setup Logging
os.makedirs("debug", exist_ok=True)
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("debug/app.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("CV-Reformatter-MVP")

app = FastAPI(title="CV Reformatter MVP")
app.state.max_file_size = 50 * 1024 * 1024  # 50MB

# CORS for Streamlit
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8501", "http://localhost:8500", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/test-upload")
def test_upload():
    return {"status": "UPLOAD OK"}


# Directories
INPUT_DIR = "input"
OUTPUT_DIR = "output"
TEMPLATES_DIR = "templates"

os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs("debug", exist_ok=True)

# Static Files
app.mount("/input", StaticFiles(directory=INPUT_DIR), name="input")
app.mount("/output", StaticFiles(directory=OUTPUT_DIR), name="output")

# Model Initialization
try:
    MODEL_NAME = "microsoft/layoutlmv3-base"
    processor = LayoutLMv3Processor.from_pretrained(MODEL_NAME, apply_ocr=True)
    model = LayoutLMv3ForSequenceClassification.from_pretrained(MODEL_NAME).to("cpu")
    logger.info("LayoutLMv3 model loaded successfully on CPU.")
except Exception as e:
    logger.error(f"Failed to load LayoutLMv3: {e}")
    processor = None
    model = None

class ProcessResponse(BaseModel):
    job_id: str
    status: str
    confidence_score: float
    extracted_data: Dict[str, Any]
    docx_url: Optional[str] = None
    certifications: Optional[List[Dict[str, str]]] = None # Not strictly needed if in extracted_data but consistent

# Robust Section Detection keywords
SECTION_KEYWORDS = [
    "summary", "profile", "experience", "education", "projects", "skills", 
    "tools", "certifications", "declaration", "languages", "personal"
]

# Hard-coded Section Boundaries
SECTION_HEADERS = [
    "profile summary", "work experience", "professional experience", 
    "education", "project details", "projects", "key skills", 
    "other skills", "tools", "certifications", "declaration", "summary", "key skills and knowledge"
]

def is_any_section_header(line: str) -> bool:
    """Strictly checks for section transitions."""
    l = line.strip().lower()
    if not l or len(l) > 40: return False
    # Exact match for common headers
    if l in SECTION_HEADERS: return True
    # Startswith match for headers that often have dates/extra text
    for k in ["education", "experience", "projects", "skills"]:
        if l.startswith(k) and len(l) < 25: return True
    return False

SECTION_HEAD_SUMMARY_RE = re.compile(r'^\s*(Profile\s+)?Summary\b|^\s*Profile\b', re.I)
SECTION_HEAD_EDUCATION_RE = re.compile(r'^\s*Education\b', re.I)
SECTION_HEAD_PROJECTS_RE = re.compile(r'^\s*Projects?|^\s*Project\s+Details\b', re.I)
SECTION_HEAD_SKILLS_RE = re.compile(r'^\s*(Key\s+)?Skills?\b|^\s*Tools\b', re.I)
SECTION_HEAD_EXPERIENCE_RE = re.compile(r'^\s*Experience|^\s*Work\s+Experience\b', re.I)
SECTION_HEAD_CERTS_RE = re.compile(r'^\s*(certifications?|courses?|awards?|training and certifications?|training)\b', re.I)
DATE_RE = re.compile(r'\b(19|20)\d{2}\b|\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b', re.I)

# Deprecated fallback
TOP_SECTION_RE = re.compile(r'^(Summary|Profile|Experience|Education|Projects|Skills|Certifications|Tools)\b', re.I)


# Preprocessing Helpers

def detect_section(line: str) -> Optional[str]:
    """Helper to detect CV sections using fuzzy matching."""
    if not line or len(line) > 50: return None
    
    # Normalize: lowercase and strip punctuation
    clean = re.sub(r'[^\w\s]', '', line.lower().strip())
    if not clean: return None

    mapping = {
        "experience": ["work experience", "professional experience", "experience", "employment history", "career history", "experience timeline", "internship experience"],
        "projects": ["projects", "project details", "academic projects", "featured projects", "major projects", "key projects"],
        "skills": ["skills", "technical skills", "key skills", "core skills", "skills & tools", "tech stack", "tools & technologies", "key skills and knowledge", "other skills", "tools", "technical skills matrix", "technology matrix", "technical competencies"]
    }

    for section, variants in mapping.items():
        for v in variants:
            # High threshold for fuzzy match to avoid false positives on list items
            if v == clean or fuzz.ratio(clean, v) > 85:
                return section
    return None

def preprocess_image(image_bytes: bytes, filename: str = "") -> Dict[str, Any]:
    """Extracts text and structural table data from PDF or Image."""
    results = {"text": "", "tables": []}
    
    if filename.lower().endswith('.pdf'):
        try:
            doc = fitz.open(stream=image_bytes, filetype="pdf")
            full_text = []
            for page in doc:
                # 1. Get raw text
                full_text.append(page.get_text())
                
                # 2. Extract tables
                tabs = page.find_tables()
                for table in tabs:
                    headers = [h.replace("\n", " ").strip() if h else f"Col{i}" for i, h in enumerate(table.header.names)]
                    table_data = table.extract() # List of lists of strings
                    
                    rows = []
                    # Start from index 1 if the first row is headers, or 0 if headers were separate
                    start_idx = 1 if table.header.external else 0 
                    # Usually table.extract() includes the header row as the first element.
                    # If table.header.names is populated, the first element of extract() is often those names.
                    
                    for r_idx, row in enumerate(table_data):
                        if r_idx == 0 and not table.header.external:
                            continue # Skip header row
                        
                        row_dict = {}
                        for i, cell in enumerate(row):
                            h = headers[i] if i < len(headers) else f"Col{i}"
                            row_dict[h] = cell.replace("\n", " ").strip() if cell else ""
                            
                        if any(row_dict.values()):
                            rows.append(row_dict)
                    
                    if rows:
                        results["tables"].append({
                            "headers": headers,
                            "rows": rows,
                            "page": page.number + 1
                        })
            
            results["text"] = "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error in preprocess_image (PDF): {e}")
            results["text"] = "Error extracting text"
    else:
        # Simple placeholder for OCR - in production we use pytesseract
        results["text"] = "OCR Text Placeholder"
        
    return results

def extract_name_and_contact(image_bytes: bytes, filename: str) -> Dict[str, str]:
    """Uses font-size heuristics to extract name and regex for contact info."""
    results = {"name": "Applicant", "email": "N/A", "phone": "N/A", "linkedin": "N/A"}
    
    if not filename.lower().endswith('.pdf'):
        return results

    try:
        doc = fitz.open(stream=image_bytes, filetype="pdf")
        first_page = doc[0]
        data = first_page.get_text("dict")
        
        spans = []
        for b in data["blocks"]:
            if "lines" in b:
                for l in b["lines"]:
                    for s in l["spans"]:
                        spans.append(s)
        
        if not spans: return results
        
        # 1. Identify NAME: All spans with the max font size in the top area
        candidate_spans = [s for s in spans if len(s['text'].strip()) > 2 and re.search('[a-zA-Z]', s['text'])]
        if candidate_spans:
            max_size = max(s['size'] for s in candidate_spans)
            # Find all spans within 1pt of max_size
            name_parts = [s['text'].strip() for s in candidate_spans if abs(s['size'] - max_size) < 1.0]
            # Join them
            full_name = " ".join(dict.fromkeys(name_parts))
            
            # GUARD: A valid name is usually 1-4 words and < 50 chars
            if len(full_name.split()) <= 4 and len(full_name) < 50:
                 results["name"] = full_name
            else:
                 # If too long, try the VERY FIRST high-confidence span only
                 first_span = candidate_spans[0]['text'].strip()
                 if len(first_span.split()) <= 4:
                     results["name"] = first_span
                 else:
                     results["name"] = "Applicant"
            
        # 2. Identify CONTACT: Use regex on the raw text of the first page
        raw_text = first_page.get_text()
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', raw_text)
        phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4,}', raw_text)
        linkedin_match = re.search(r'linkedin\.com/in/[\w\.-]+', raw_text)
        
        if email_match: results["email"] = email_match.group(0)
        if phone_match: results["phone"] = phone_match.group(0)
        if linkedin_match: results["linkedin"] = linkedin_match.group(0)
            
    except Exception as e:
        logger.error(f"Error extracting name: {e}")
        
    return results

def is_responsibility_line(line: str) -> bool:
    """Detects if a line represents an action or responsibility."""
    # Strip bullets and symbols
    clean = line.lstrip("▪•- *").strip()
    if not clean: return False
    
    # Rule 1: Starts with a bullet point
    if any(line.lstrip().startswith(b) for b in ["▪", "•", "-", "*"]):
        return True
    
    # Rule 2: Starts with a strong action verb
    verbs = ["designed", "developed", "implemented", "built", "led", "managed", 
             "created", "optimized", "improved", "migrated", "architected", "configured"]
    
    first_word = re.sub(r'[^\w]', '', clean.split()[0].lower()) if clean.split() else ""
    if first_word in verbs:
        return True
        
    return False

def is_location(text: str) -> bool:
    """Heuristic to detect if a string is likely a location."""
    if not text: return False
    places = ["kochi", "bangalore", "mumbai", "india", "trivandrum", "ireland", "us", "thailand", "san francisco", "ca", "calicut"]
    text_lower = text.lower()
    
    # Check if any place is a standalone word or the string ends with it
    for p in places:
        if p in text_lower and len(text.split()) <= 3:
            return True
    if text_lower.endswith(", india") or text_lower.endswith(", us"):
        return True
    return False

def get_cv_segments(input_data: Any) -> Dict[str, Dict[str, Any]]:
    """Splits raw CV text and tables into logical section blocks."""
    if isinstance(input_data, dict):
        text = input_data.get("text", "")
        raw_tables = input_data.get("tables", [])
    else:
        text = str(input_data)
        raw_tables = []

    lines = text.split('\n')
    segments = {
        "summary": {"text": [], "tables": []},
        "experience": {"text": [], "tables": []},
        "education": {"text": [], "tables": []},
        "projects": {"text": [], "tables": []},
        "skills": {"text": [], "tables": []},
        "certifications": {"text": [], "tables": []},
        "meta": {"text": [], "tables": []}
    }
    
    current_key = "summary"
    
    # Section Header Detection Patterns
    SECT_MAP = {
        "experience": ["work experience", "professional experience", "employment history", "career history", "experience summary", "employment details"],
        "education": ["education", "academic qualification", "academic profile", "academic background", "academics", "educational qualification"],
        "projects": ["projects", "key projects", "project details", "project portfolio", "technical projects", "major projects", "project profile"],
        "skills": ["skills", "technical skills", "key skills", "skill set", "core competencies", "tools", "technologies", "tech stack", "software skills", "other skills", "key skills and knowledge"],
        "certifications": ["certifications", "training", "courses", "awards", "achievement"],
        "summary": ["profile summary", "professional summary", "summary", "profile", "profile highlights", "career highlights", "professional profile"],
        "experience_junk": ["chipsnbytes", "technolog", "developer", "engineer", "designer"]
    }

    # 1. Segment Text
    for line in lines:
        l_orig = line.strip()
        if not l_orig: continue
        
        l_clean = re.sub(r'^[•▪\-\*▪➢\d\.\s\t]+', '', l_orig).strip().lower()
        
        found_new = False
        if 2 < len(l_clean) < 45:
             for key, synonyms in SECT_MAP.items():
                  # A label like 'Tools:' in projects should NOT switch the section back to skills
                  if key == "skills" and current_key in ["projects", "experience"] and ":" in l_orig:
                      continue
                      
                  if any(l_clean == s or l_clean.startswith(s + " ") or l_clean.startswith(s + ":") or (s in l_clean and len(l_clean) < len(s) + 5) for s in synonyms):
                      if key == "experience_junk" and current_key == "skills":
                          current_key = "meta"
                          found_new = True
                          break
                      if key != current_key and key != "experience_junk":
                          logger.info(f"Segment Switch: {current_key} -> {key} at line: '{l_orig}'")
                          current_key = key
                          found_new = True
                          break
                      if key == current_key:
                          found_new = True
                          break
        
        if not found_new:
            segments[current_key]["text"].append(line)
            
    # Join text lists
    for k in segments:
        segments[k]["text"] = "\n".join(segments[k]["text"])

    # 2. Associate Tables with Sections
    # Purely heuristic: If a table contains keywords for a section, put it there.
    for table in raw_tables:
        table_str = json.dumps(table).lower()
        matched = False
        for key, synonyms in SECT_MAP.items():
            if any(s in table_str for s in synonyms if len(s) > 4):
                if key != "experience_junk":
                    segments[key]["tables"].append(table)
                    matched = True
                    break
        
        if not matched:
            # Fallback: Put in current section if it was experience or projects
            if "project" in table_str: segments["projects"]["tables"].append(table)
            elif "experience" in table_str or "employer" in table_str: segments["experience"]["tables"].append(table)
            else: segments["experience"]["tables"].append(table) # Default to experience for structured data

    return segments

def extract_summary(input_data: Any) -> str:
    """Standard summary extraction - Stable Version."""
    text = input_data["text"] if isinstance(input_data, dict) else str(input_data)
    
    noise_re = re.compile(r'^[•▪\-\*▪\t\s]+')
    sentences = []
    
    text = text.replace("[FILL HERE]", "").strip()
    for line in text.split('\n'):
        l = noise_re.sub('', line).strip()
        if not l or len(l) < 10: continue
        
        candidates = re.split(r'(?<=[.!?])\s+', l)
        for cand in candidates:
            cand = cand.strip()
            if not cand: continue
            if not any(fuzz.ratio(cand.lower(), ex.lower()) > 85 for ex in sentences):
                sentences.append(cand)
                
    result = " ".join(sentences)
    if result and result[-1].isalnum():
        result += "."
    return result

def extract_skills_it(input_data: Any, candidate_name: str = "") -> Dict[str, List[str]]:
    """Industry-standard technical skill and tool extraction."""
    if isinstance(input_data, dict):
        text = input_data.get("text", "")
        tables = input_data.get("tables", [])
    else:
        text = str(input_data)
        tables = []

    tech_skills = []
    tech_tools = []
    noise_re = re.compile(r'^[•▪\-\*▪\x00-\x1f\x7f-\x9f\s\t/]+')
    name_parts = candidate_name.lower().split() if candidate_name else []
    
    # 1. Core Technical Knowledge Base (Languages, Frameworks, Methods)
    core_tech_dict = [
        "php", "laravel", "yii", "codeigniter", "wordpress", "magento", "symfony",
        "javascript", "typescript", "react", "angular", "vue", "next.js", "nuxtjs", "lit", "litelement",
        "node", "express", "python", "django", "flask", "fastapi", "java", "spring", "c#", "asp.net",
        "html", "css", "html5", "css3", "sass", "less", "tailwind", "bootstrap",
        "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "oracle",
        "ui/ux", "ui design", "ux design", "user research", "wireframing", "prototyping",
        "graphic design", "branding", "interaction design", "sitemaps", "user flows",
        "agile", "scrum", "sdlc", "rest api", "json", "ajax", "web sockets", "microservices"
    ]
    
    # 2. Industry Tools Base (Software, Services, Hardware)
    industry_tools_dict = [
        "figma", "adobe xd", "sketch", "invision", "photoshop", "illustrator", "figjam", "zeplin",
        "git", "github", "gitlab", "bitbucket", "docker", "kubernetes", "aws", "azure", "gcp",
        "postman", "jira", "asana", "confluence", "slack", "vscode", "sublime", "cpanel", "draw.io",
        "kibana", "grafana", "jenkins", "terraform", "maven", "npm", "yarn"
    ]
    
    # Junk to discard (Soft skills or non-IT fluff)
    junk_filters = ["communication", "team work", "interpersonal", "fast learning", "highly adaptive", "flexible", "team player", "problem solving", "other skills", "knowledge"]

    # Collect all items to process (from text lines and table cells)
    source_items = text.split('\n')
    for table in tables:
        for row in table.get("rows", []):
            source_items.extend(row.values())

    for line in source_items:
        line = str(line).strip()
        if not line or len(line) > 1000: continue
        if name_parts and any(p in line.lower() for p in name_parts if len(p) > 2): continue
        
        line = noise_re.sub('', line).strip()
        # Handle both comma/bullet separated and space-separated if long
        parts = re.split(r'[,\u2022•;|]', line)
        for part in parts:
            t = part.strip(" .()[]/\\")
            if not t or len(t) < 2: continue
            
            t_low = t.lower()
            if any(j in t_low for j in junk_filters): continue

            # Check for location or section header bleed
            if is_location(t) or t_low in SECTION_HEADERS:
                continue

            # Match against tech/tools dictionaries to avoid noise
            is_tech = False
            for c in core_tech_dict:
                if re.search(r'\b' + re.escape(c) + r'\b', t_low):
                    # For multi-word matches like 'ui/ux', use the dictionary form
                    val = c.upper() if len(c) < 5 else c.title()
                    # Custom casing
                    if c == "ui/ux": val = "UI/UX"
                    if c == "mysql": val = "MySQL"
                    if c == "php": val = "PHP"
                    if val not in tech_skills: tech_skills.append(val)
                    is_tech = True
                    # Don't break, a line might contain multiple known skills
            
            is_tool = False
            for tool in industry_tools_dict:
                if re.search(r'\b' + re.escape(tool) + r'\b', t_low):
                    val = tool.upper() if len(tool) < 4 else tool.title()
                    if val not in tech_tools: tech_tools.append(val)
                    is_tool = True
            
            # Fallback: If it wasn't a known "Tech" or "Tool", but it appeared in this section...
            # and it's short and not junk, treat it as a generic skill.
            if not is_tech and not is_tool and len(t) < 30:
                if t not in tech_skills:
                    tech_skills.append(t)
                    is_tool = True
            
            if not is_tech and not is_tool:
                # Fallback for acronyms or capitalized products not in our small dicts
                if len(t) <= 15 and any(c.isupper() for c in t) and not is_location(t):
                     t_low = t.lower()
                     if t_low not in tech_skills and t_low not in ["other", "skills", "knowledge", "role", "duration", "tech stack", "project", "details", "description"]:
                        # Strict check: ignore patterns like 'Project #1', 'Project Title', 'Duration:'
                        if "project" in t_low or "title" in t_low or "duration" in t_low or "#" in t_low:
                            continue
                        tech_skills.append(t)

    return {"skills": tech_skills, "tools": tech_tools}

    return {"skills": tech_skills, "tools": tech_tools}

# Alias for backward compatibility
extract_skills = extract_skills_it

def is_valid_institution(text: str) -> bool:
    """Strict institution validation."""
    if not text or len(text) < 5: return False
    l_low = text.lower()
    
    # Aggressive block list
    block_words = ["training", "certification", "experience", "summary", "project", "details", "highlights", "profile", "career", "technical"]
    if any(x in l_low for x in block_words):
        return False
        
    # Valid indicators
    inst_keywords = ["university", "college", "institute", "school", "academy", "vidyalaya", "management", "science", "technology"]
    if any(k in l_low for k in inst_keywords):
        return True
    
    # Length and casing check
    words = text.split()
    if 2 <= len(words) <= 8 and all(w[0].isupper() for w in words if w.isalpha()):
        return True
        
    return False

def is_likely_role(text: str) -> bool:
    """Check if text contains role-related keywords."""
    roles = ["engineer", "developer", "manager", "lead", "architect", "analyst", "coder", "designer", "specialist", "intern", "trainee"]
    text_lower = text.lower()
    return any(r in text_lower for r in roles)

def extract_work_experience(input_data: Any) -> List[Dict[str, str]]:
    """Standard experience extraction - Stable Phase 2 Version."""
    if isinstance(input_data, dict):
        text = input_data.get("text", "")
        tables = input_data.get("tables", [])
    else:
        text = str(input_data)
        tables = []

    experience = []

    # 1. Structural Table Extraction (High Fidelity)
    for table in tables:
        headers = [h.lower() for h in table.get("headers", [])]
        rows = table.get("rows", [])
        
        # Check if this table looks like an experience grid
        exp_markers = ["company", "employer", "organization", "period", "duration", "role", "designation"]
        if any(m in " ".join(headers) for m in exp_markers):
            for row in rows:
                item = {"company": "N/A", "role": "N/A", "duration": "N/A", "location": "N/A", "responsibilities": ""}
                for k, v in row.items():
                    kl = k.lower()
                    if "company" in kl or "employer" in kl: item["company"] = v
                    elif "role" in kl or "designation" in kl: item["role"] = v
                    elif "duration" in kl or "period" in kl: item["duration"] = v
                    elif "location" in kl: item["location"] = v
                    elif "responsibilities" in kl or "description" in kl or "details" in kl: item["responsibilities"] = v
                
                if item["company"] != "N/A" or item["role"] != "N/A":
                    experience.append(item)

    # 2. Legacy Regex-based Extraction
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
            
        has_date = bool(DATE_RE.search(line))
        has_sep = any(s in line for s in ['|', '-', '–', ',', '\t'])
        
        if has_date and has_sep and len(line) < 150:
            if not is_responsibility_line(line):
                # Split using pipes or dashes first, as they are cleaner field separators
                parts = [p.strip() for p in re.split(r'\s*[|–-]\s*|\t+', line) if p.strip()]
                
                # Intelligent Field Assignment
                company = parts[0]
                role = "N/A"
                duration = "N/A"
                location = "N/A"
                
                # If we have multiple parts, try to identify which is which
                other_parts = parts[1:]
                for p in other_parts:
                    if bool(DATE_RE.search(p)):
                        duration = p
                    elif is_likely_role(p):
                        role = p
                    elif is_location(p) or "india" in p.lower():
                        location = p
                
                # If role is still N/A, take the first non-date, non-location part
                if role == "N/A":
                    for p in other_parts:
                        if p != duration and p != location:
                            role = p; break

                # Fallback for comma-separated single line: "Company, Role, Date"
                if role == "N/A" and ',' in company:
                    c_parts = [cp.strip() for cp in company.split(',')]
                    if len(c_parts) >= 2:
                        company = c_parts[0]
                        role = c_parts[1]
                        if len(c_parts) > 2 and duration == "N/A":
                            duration = c_parts[-1]

                # Role Lookahead
                if role == "N/A" and i + 1 < len(lines):
                    next_l = lines[i+1].strip()
                    if next_l and not is_responsibility_line(next_l) and not DATE_RE.search(next_l):
                        role = next_l
                        i += 1
                
                resps = []
                j = i + 1
                while j < len(lines):
                    l_sub = lines[j].strip()
                    if not l_sub or is_any_section_header(l_sub): break
                    
                    # Stop if a new job header is detected
                    if bool(DATE_RE.search(l_sub)) and any(s in l_sub for s in ['|', '-', '–', '\t', ',']):
                        break
                        
                    if is_responsibility_line(l_sub) or len(l_sub) > 20:
                        resps.append(l_sub.lstrip("▪•- *").strip())
                    j += 1
                
                # Check for duplicates before adding
                is_dup = False
                for existing in experience:
                    if fuzz.ratio(company.lower(), existing["company"].lower()) > 85 and \
                       fuzz.ratio(role.lower(), existing["role"].lower()) > 85:
                        is_dup = True
                        break
                
                if not is_dup:
                    experience.append({
                        "company": company,
                        "role": role,
                        "duration": duration,
                        "location": location,
                        "responsibilities": " ".join(resps)
                    })
                i = j
                continue
        i += 1
    return experience

def extract_projects(input_data: Any, companies: List[str] = None) -> List[Dict[str, Any]]:
    """Identifies projects while filtering out company names from experience."""
    if isinstance(input_data, dict):
        text = input_data.get("text", "")
        tables = input_data.get("tables", [])
    else:
        text = str(input_data)
        tables = []

    projects = []
    
    # 1. Structural Table Extraction
    for table in tables:
        headers = [h.lower() for h in table.get("headers", [])]
        rows = table.get("rows", [])
        
        proj_markers = ["project", "client", "tech", "role", "description"]
        if any(m in " ".join(headers) for m in proj_markers):
            for row in rows:
                item = {"title": "Project", "tech": "N/A", "duration": "N/A", "role": "N/A", "details": ""}
                for k, v in row.items():
                    kl = k.lower()
                    if "title" in kl or "project" in kl: item["title"] = v
                    elif "stack" in kl or "technolog" in kl: item["tech"] = v
                    elif "role" in kl: item["role"] = v
                    elif "duration" in kl or "period" in kl: item["duration"] = v
                    elif "description" in kl or "details" in kl or "summary" in kl: item["details"] = v
                    elif "responsibilities" in kl: item["responsibilities"] = v
                
                # Pradeep's CV specific cleanup: If title is empty but we have 'Project #1:', fix it
                if item["title"].lower().startswith("project") and not v and ":" in k:
                     # This handles cases where the key is "Project #1:" and value is empty
                     item["title"] = k.rstrip(":")

                if item["title"] != "Project" or item["details"] or item["tech"] != "N/A":
                    projects.append(item)

    # 2. Legacy Extraction
    lines = text.split('\n')
    current = None
    
    company_names = [c.lower() for c in (companies or [])]
    
    def start_new(title):
        return {"title": title, "tech": "N/A", "duration": "N/A", "role": "N/A", "desc": [], "role_text": ""}

    def flush_project():
        nonlocal current
        if current:
            # Validate: Title shouldn't be a company name or section header
            t_low = current['title'].lower()
            if any(c in t_low for c in company_names) and len(t_low) < 50:
                 current = None
                 return
            
            if len(current['title']) < 3 or is_any_section_header(current['title']):
                current = None
                return

            # Check for duplicates from tables
            is_dup = False
            for p in projects:
                if fuzz.ratio(current['title'].lower(), p['title'].lower()) > 90:
                    is_dup = True
                    break
            if is_dup:
                current = None
                return

            desc_str = " ".join(current['desc'])
            # Basic cleanup of desc
            desc_str = desc_str.replace("Responsibilities:", "").strip()
            
            projects.append({
                "title": current['title'],
                "tech": current['tech'] if current.get('tech') != "N/A" else current.get('stack', 'N/A'),
                "duration": current['duration'],
                "role": current.get('role_text', current['role']),
                "responsibilities": current.get('responsibilities', ""),
                "details": desc_str
            })
            current = None

    # Metadata patterns
    META_KEYS = {
        "client": re.compile(r'^(client|customer|organization)\s*:', re.I),
        "role": re.compile(r'^(role|position)\s*:', re.I),
        "tech": re.compile(r'^(tech stack|technologies|environment|tools|stack)\s*:', re.I),
        "duration": re.compile(r'^(duration|period|time)\s*:', re.I),
        "responsibilities": re.compile(r'^(responsibilities|tasks|duties|roles and responsibilities)\s*:', re.I),
        "overview": re.compile(r'^(project overview|description|summary|details)\s*:', re.I)
    }

    for i, line in enumerate(lines):
        l = re.sub(r'^[•▪\-\*▪\x00-\x1f\x7f-\x9f\s]+', '', line).strip()
        if not l: continue
        l_low = l.lower()
        
        # Meta Check
        is_meta = False
        for key, ptrn in META_KEYS.items():
            if ptrn.match(l):
                is_meta = True
                val = l.split(":", 1)[1].strip()
                if not current: current = start_new("Project")
                
                if key == "tech": current['tech'] = val
                elif key == "duration": current['duration'] = val
                elif key == "role": current['role_text'] = val
                elif key == "responsibilities": 
                    current['responsibilities'] = val
                    if not current['desc']: current['desc'].append(val) # Fallback to desc
                elif key == "overview": current['desc'].append(val)
                break
        
        if is_meta: continue
        
        # Piped Header Check (Role | Project | Date)
        if '|' in l and len(l) < 120:
            flush_project()
            parts = l.split('|')
            current = start_new(parts[0].strip())
            if len(parts) > 1: 
                p2 = parts[1].strip()
                # Heuristic: If it contains tech keywords, it's tech, not role
                if any(k in p2.lower() for k in ["code", "git", "php", "aws", "stack", "env", "sql", "js"]):
                    current['tech'] = p2
                else:
                    current['role_text'] = p2
            continue
            
        # Title detection
        if 5 < len(l) < 100 and not l.endswith('.') and l[0].isupper():
             # If next line looks like project meta, it's definitely a title
             is_def_title = False
             if i + 1 < len(lines):
                  nl = lines[i+1].strip().lower()
                  if any(k in nl for k in ["role:", "tech stack:", "technologies:", "project overview:"]):
                      is_def_title = True
             
             if is_def_title:
                 flush_project()
                 current = start_new(l)
                 continue

        if current:
            if is_any_section_header(l):
                flush_project()
            else:
                if l not in current['desc']:
                    current['desc'].append(l)

    flush_project()
    return projects

def extract_education(input_data: Any) -> List[Dict[str, str]]:
    """Strict education extraction - requires both degree keyword AND reasonable context."""
    if isinstance(input_data, dict):
        text = input_data.get("text", "")
        tables = input_data.get("tables", [])
    else:
        text = str(input_data)
        tables = []

    education = []

    # 1. Table Extraction
    for table in tables:
        headers = [h.lower() for h in table.get("headers", [])]
        rows = table.get("rows", [])
        
        edu_markers = ["degree", "qualification", "university", "college", "institute", "year", "pass"]
        if any(m in " ".join(headers) for m in edu_markers):
            for row in rows:
                item = {"degree": "N/A", "institution": "N/A", "duration": "N/A"}
                for k, v in row.items():
                    kl = k.lower()
                    if "degree" in kl or "qualification" in kl: item["degree"] = v
                    elif "university" in kl or "college" in kl or "institute" in kl: item["institution"] = v
                    elif "year" in kl or "pass" in kl or "duration" in kl: item["duration"] = v
                
                if item["degree"] != "N/A" or item["institution"] != "N/A":
                    education.append(item)

    # 2. Legacy Regex Logic
    lines = text.split('\n')
    degree_keywords = ["bachelor", "master", "degree", "diploma", "b.a", "b.s", "m.a", "m.s", "btech", "mtech", "b.tech", "m.tech", "bca", "mca", "b.sc", "m.sc", "graduate", "certificate"]
    
    for i, line in enumerate(lines):
        l = re.sub(r'^[•▪\-\*▪\s\t]+', '', line).strip()
        if not l or len(l) < 5: continue
        l_low = l.lower()
        
        if not any(k in l_low for k in degree_keywords): continue
        if "|" in l or " - " in l:
            if any(r in l_low for r in ["engineer", "developer", "manager", "lead", "architect", "analyst", "technologies", "informatics"]):
                continue

        if len(l.split()) > 15: continue

        year_match = re.search(r'\b(19|20)\d{2}\b', l)
        year = year_match.group(0) if year_match else "N/A"
        
        title = l.replace(year, "").strip(" -–,|•▪")
        parts = re.split(r'[–-]|\|', title)
        degree = parts[0].strip()
        inst = "N/A"
        
        if len(parts) > 1:
            inst = parts[1].strip()
        else:
            for k in range(max(0, i-1), min(i+3, len(lines))):
                if k == i: continue
                nl = lines[k].strip()
                if nl and is_valid_institution(nl) and not any(kw in nl.lower() for kw in degree_keywords):
                    inst = nl
                    break
        
        if degree:
            # Deduplicate
            is_dup = False
            for existing in education:
                if fuzz.partial_ratio(degree.lower(), existing["degree"].lower()) > 85:
                    is_dup = True; break
            if not is_dup:
                education.append({"degree": degree, "institution": inst, "duration": year})

    return education

def extract_certifications(input_data: Any) -> List[Dict[str, str]]:
    """Generic certification extraction handling wrapped lines/URLs."""
    if isinstance(input_data, dict):
        text = input_data.get("text", "")
        tables = input_data.get("tables", [])
    else:
        text = str(input_data)
        tables = []

    cert_lines = []
    
    # Process tables if they contain certifications
    for table in tables:
        headers = " ".join(table.get("headers", [])).lower()
        if "cert" in headers or "course" in headers or "award" in headers:
            for row in table.get("rows", []):
                cert_lines.append(" - ".join([str(v) for v in row.values() if v]))

    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        is_bullet = any(line.startswith(b) for b in ["•", "-", "*"])
        if is_bullet:
            cert_lines.append(line.lstrip("•- *").strip())
        elif line.lower().startswith("http"):
            if cert_lines: cert_lines[-1] += " " + line
            else: cert_lines.append(line)
        else:
            if cert_lines: cert_lines[-1] += " " + line
            else: cert_lines.append(line)
                 
    certs = []
    for raw in cert_lines:
        url = "N/A"
        url_match = re.search(r'https?://[^\s,]+', raw)
        title = raw
        issuer = "N/A"
        
        if url_match:
            url = url_match.group(0)
            title = raw.replace(url, "").strip(" -–,")
            
        if '-' in title or '–' in title:
             parts = re.split(r'[–-]', title, maxsplit=1)
             if len(parts) >= 2:
                 title = parts[0].strip()
                 issuer = parts[1].strip()
        
        if title:
            certs.append({"title": title, "issuer": issuer, "url": url})
        
    return certs

def generate_docx(data: Dict[str, Any], template_path: str, output_path: str):
    """Builds a premium, modern standard CV with invisible table layouts."""
    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Header Section: Name & Contact
        name = data.get("full_name", "Applicant")
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(41, 128, 185) # Modern Blue
        
        contact_line = []
        if data.get("email") and data.get("email") != "N/A": contact_line.append(data["email"])
        if data.get("phone") and data.get("phone") != "N/A": contact_line.append(data["phone"])
        if data.get("linkedin") and data.get("linkedin") != "N/A": contact_line.append(data["linkedin"])
        
        contact_p = doc.add_paragraph(" | ".join(contact_line))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(20)

        # 1. Summary Section
        if data.get("summary"):
            doc.add_heading('PROFESSIONAL SUMMARY', level=1)
            doc.add_paragraph(data["summary"])
        
        # 2. Skills Section
        if data.get("skills"):
            doc.add_heading('TECHNICAL SKILLS', level=1)
            skills_str = ", ".join(data.get("skills", []))
            doc.add_paragraph(skills_str)
        
        # 3. Experience Section - Invisible Table
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        for job in data.get("work_experience", []):
            exp_table = doc.add_table(rows=1, cols=2)
            exp_table.autofit = True
            
            # Left Column: Role & Company
            role_cell = exp_table.rows[0].cells[0]
            r_para = role_cell.paragraphs[0]
            r_run = r_para.add_run(f"{job.get('role', 'Role')}\n")
            r_run.bold = True
            r_para.add_run(job.get('company', 'Company'))
            
            # Right Column: Dates & Location
            date_cell = exp_table.rows[0].cells[1]
            date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_cell.text = f"{job.get('duration', 'N/A')}\n{job.get('location', '')}"
            
            # Responsibilities as bullets
            resps = job.get("responsibilities", "") or job.get("resps", "")
            if resps:
                # Basic cleaning of the text before splitting
                clean_resps = resps.replace("[FILL HERE]", "").strip()
                for sentence in re.split(r'(?<=[\.\!\?])\s+', clean_resps):
                    if len(sentence.strip()) > 10:
                        p = doc.add_paragraph(style='List Bullet')
                        p.text = sentence.strip()
                        p.paragraph_format.space_after = Pt(2)
            
            doc.add_paragraph() # Spacer

        # 4. Projects Section
        if data.get("projects"):
            doc.add_heading('PROJECT PORTFOLIO', level=1)
            for item in data.get("projects", []):
                p_head = doc.add_paragraph()
                p_run = p_head.add_run(f"{item.get('title', 'Project')} | {item.get('tech', 'N/A')}")
                p_run.bold = True
                
                # Role and Details
                if item.get("role") and item.get("role") != "N/A":
                    doc.add_paragraph(f"Role: {item['role']}")
                
                details = item.get("details", "")
                if details:
                    # Clean the details up
                    clean_details = details.replace("[FILL HERE]", "").strip()
                    for sentence in re.split(r'(?<=[\.\!\?])\s+', clean_details):
                        if len(sentence.strip()) > 10:
                            p = doc.add_paragraph(style='List Bullet')
                            p.text = sentence.strip()
                            p.paragraph_format.space_after = Pt(2)
                doc.add_paragraph()
            
        # 5. Education Section
        if data.get("education"):
            doc.add_heading('EDUCATION', level=1)
            for edu in data.get("education", []):
                p = doc.add_paragraph()
                e_run = p.add_run(f"{edu.get('degree', 'Degree')} ")
                e_run.bold = True
                p.add_run(f"- {edu.get('institution', 'N/A')} ({edu.get('duration', 'N/A')})")
        
        # Clean up borders for all tables (Invisible Table Theme)
        for table in doc.tables:
            table.style = 'Normal Table'
            
        # Versioning/Success Stamp in Footer
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = "Engineered by CV Reformatter v1.1 | Final Processed Layout"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.save(output_path)
        logger.info(f"Generated Premium Standard CV at {output_path}")
    except Exception as e:
        logger.error(f"Premium DOCX Generation Error: {e}")
        raise

@app.get("/templates/list", response_model=List[str])
def api_list_templates():
    """List all registered templates."""
    try:
        return list_templates()
    except Exception as e:
        logger.error(f"Template list error: {e}")
        return []

from template_engine.template_cleaner import clean_template_content

@app.post("/templates/upload")
async def upload_template_endpoint(file: UploadFile = File(...), template_name: str = "New Template"):
    """Upload a DOCX template, extract schema, CLEAN it, and register."""
    try:
        content = await file.read()
        temp_path = os.path.join(TEMPLATES_DIR, file.filename)
        
        # 1. Save Original Temporarily
        with open(temp_path, "wb") as f:
            f.write(content)
            
        # 2. Open Document Instance ONCE (to keep object IDs consistent)
        doc = Document(temp_path)
        
        # 3. Extract Schema
        schema = extract_template_schema(temp_path, template_name, doc=doc)
        
        # 4. Clean Template (REMOVE DUPLICATES)
        doc = clean_template_content(doc, schema)
        doc.save(temp_path)
        
        # 5. Register
        register_template(temp_path, template_name)
        
        return {"status": "success", "template_name": template_name, "sections_found": len(schema.sections), "schema": schema.dict()}
    except Exception as e:
        logger.error(f"Template upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/process-to-template")
async def process_cv_to_template(file: UploadFile = File(...), template_name: str = "default"):
    """Extract CV data and fill ONLY the selected template."""
    job_id = str(uuid.uuid4())
    try:
        # 1. Extract CV Data (Reuse Phase 1 Logic)
        content = await file.read()
        text = preprocess_image(content, file.filename)
        
        # 0. Personal Info Extraction (New)
        personal = extract_name_and_contact(content, file.filename)
        
        # 1. Segmented Extraction
        segs = get_cv_segments(text)
        skills_obj = extract_skills_it(segs.get("skills", ""), personal.get("name", "Applicant"))
        
        # Collect company names to filter projects
        work_exp = extract_work_experience(segs.get("experience", ""))
        companies = [job.get("company", "") for job in work_exp if job.get("company")]

        extracted = {
            "full_name": personal.get("name", "Applicant"),
            "email": personal.get("email", "N/A"),
            "phone": personal.get("phone", "N/A"),
            "linkedin": personal.get("linkedin", "N/A"),
            "summary": extract_summary(segs.get("summary", "")),
            "skills": skills_obj["skills"],
            "tools": skills_obj["tools"],
            "work_experience": work_exp,
            "projects": extract_projects(segs.get("projects", ""), companies),
            "education": extract_education(segs.get("education", "")),
            "certifications": extract_certifications(segs.get("certifications", ""))
        }
        
        # 2. Load Template
        if template_name == "default":
             # Fallback to hardcoded Phase 1 generation if default? 
             # No, let's use the new system if possible, or just call generate_docx
             pass
        
        schema = get_template_schema(template_name)
        if not schema:
            raise HTTPException(status_code=404, detail="Template not found")
            
        # 3. Fill Template
        output_filename = f"{job_id}_{template_name}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        fill_template(schema, extracted, output_path)
        
        return FileResponse(output_path, filename=output_filename)

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.exception(f"Template process error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/process", response_model=ProcessResponse)
async def process_cv(file: UploadFile = File(...)):
    """
    Standard Processing (Phase 1) - UPGRADED to use Phase 2 Engine
    Now acts as 'Process to Default Template' using Extractor_Master.
    """
    job_id = str(uuid.uuid4())
    logger.info(f"Processing Upload: {file.filename} ({job_id})")

    try:
        content = await file.read()
        
        # 1. Preprocess & Extract
        pre_res = preprocess_image(content, file.filename)
        text = pre_res.get("text", "")
        
        # Personal Info
        personal = extract_name_and_contact(content, file.filename)
        
        # Segmented Extraction
        segs = get_cv_segments(pre_res) # Pass pre_res (dict) not text str
        skills_obj = extract_skills_it(segs.get("skills", ""), personal.get("name", "Applicant"))
        
        work_exp = extract_work_experience(segs.get("experience", ""))
        companies = [job.get("company", "") for job in work_exp if job.get("company")]

        extracted = {
            "full_name": personal.get("name", "Applicant"),
            "email": personal.get("email", "N/A"),
            "phone": personal.get("phone", "N/A"),
            "linkedin": personal.get("linkedin", "N/A"),
            "summary": extract_summary(segs.get("summary", "")),
            "skills": skills_obj["skills"],
            "tools": skills_obj["tools"],
            "work_experience": work_exp,
            "projects": extract_projects(segs.get("projects", ""), companies),
            "education": extract_education(segs.get("education", "")),
            "certifications": extract_certifications(segs.get("certifications", ""))
        }

        # 2. Use Extractor_Master Template (Phase 2 Engine)
        # Ensure the template exists first
        template_name = "Extractor_Master"
        schema = get_template_schema(template_name)
        
        output_filename = f"{job_id}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        if schema:
            fill_template(schema, extracted, output_path)
            logger.info("Using Extractor_Master template for Standard Process.")
        else:
            # Fallback (Should not happen if setup correctly)
            logger.warning("Extractor_Master not found. Creating empty doc.")
            Document().save(output_path)

        return ProcessResponse(
            job_id=job_id,
            status="success",
            confidence_score=0.95,
            extracted_data=extracted,
            docx_url=f"/output/{output_filename}"
        )
    except Exception as e:
        logger.exception(f"Process error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/status")
async def get_status():
    p = len([f for f in os.listdir(INPUT_DIR) if f.endswith(".json")])
    f = len([f for f in os.listdir(OUTPUT_DIR) if f.endswith(".docx")])
    return {"pending": p, "formatted": f}

# PHASE 3: ISOLATED JD OPTIMIZER ENDPOINTS
@app.post("/v3/fill-template")
async def v3_fill_template_endpoint(data: Dict[str, Any], template_name: str = "default"):
    """
    PHASE 3: Optimized Fill using ISOLATED core logic.
    Ensures Phase 2 logic is NEVER touched.
    """
    from jd_optimizer.core.template_mapper import fill_template as v3_fill
    from main import get_template_schema, OUTPUT_DIR
    import uuid
    
    job_id = str(uuid.uuid4())
    try:
        schema = get_template_schema(template_name)
        if not schema:
            raise HTTPException(status_code=404, detail="Template not found")
            
        output_filename = f"V3_{job_id}_{template_name}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        # Use ISOLATED mapper
        v3_fill(schema, data, output_path)
        
        return FileResponse(output_path, filename=output_filename)
    except Exception as e:
        logger.exception(f"V3 process error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
